"""Export a meeting transcript for sharing: Word (.docx), PDF, or plain text.

Files land in ~/Downloads (the natural hand-off spot) and are revealed in
Finder. PDF renders a styled HTML through the locally-installed Chrome in
headless mode — no extra Python dependencies. Everything stays on-device.
"""
import html
import json
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path

from . import config, dates

CHROME = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"


def _load(base: str) -> dict:
    return json.loads(config.meeting_file(base, ".json").read_text())


def _fmt_ts(sec: float) -> str:
    sec = int(round(sec))
    h, m, s = sec // 3600, (sec % 3600) // 60, sec % 60
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _meta_line(base: str, d: dict) -> str:
    parts = []
    iso = dates.meeting_date(base)
    if iso:
        parts.append(datetime.fromisoformat(iso).strftime("%B %-d, %Y"))
    parts.append(f"{round(d.get('duration_sec', 0) / 60)} min")
    who = [s["display"] for s in d.get("speakers", [])]
    if who:
        parts.append(", ".join(who))
    return "  ·  ".join(parts)


def _out_path(base: str, ext: str) -> Path:
    out = Path.home() / "Downloads" / f"{base}.{ext}"
    n = 1
    while out.exists():
        out = Path.home() / "Downloads" / f"{base} ({n}).{ext}"
        n += 1
    return out


def to_docx(base: str) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor

    d = _load(base)
    doc = Document()
    doc.add_heading(base, level=1)
    meta = doc.add_paragraph(_meta_line(base, d))
    meta.runs[0].font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)
    meta.runs[0].font.size = Pt(10)

    for seg in d.get("segments", []):
        text = seg.get("text", "").strip()
        if not text:
            continue
        who = seg.get("display") or seg.get("name") or seg.get("speaker") or "?"
        p = doc.add_paragraph()
        r = p.add_run(f"{who}  [{_fmt_ts(seg['start'])}]"
                      + ("  ⚠" if seg.get("flags") else ""))
        r.bold = True
        r.font.size = Pt(10)
        doc.add_paragraph(text)
    out = _out_path(base, "docx")
    doc.save(str(out))
    return out


def to_html(base: str) -> str:
    d = _load(base)
    rows = []
    for seg in d.get("segments", []):
        text = seg.get("text", "").strip()
        if not text:
            continue
        who = seg.get("display") or seg.get("name") or seg.get("speaker") or "?"
        warn = ' <span class="warn">uncertain</span>' if seg.get("flags") else ""
        rows.append(
            f'<div class="seg"><div class="who">{html.escape(who)}'
            f' <span class="ts">{_fmt_ts(seg["start"])}</span>{warn}</div>'
            f'<div class="txt">{html.escape(text)}</div></div>')
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>
body{{font:11pt -apple-system,'Helvetica Neue',sans-serif;color:#1d1d1f;
margin:48px;max-width:680px}}
h1{{font-size:17pt;margin:0 0 4px}}
.meta{{color:#6e6e73;font-size:9pt;margin-bottom:22px}}
.seg{{margin-bottom:13px;page-break-inside:avoid}}
.who{{font-weight:600;font-size:9pt}}
.ts{{color:#86868b;font-weight:400}}
.warn{{color:#c93400;font-weight:400;font-size:8pt}}
.txt{{margin-top:1px}}
</style></head><body><h1>{html.escape(base)}</h1>
<div class="meta">{html.escape(_meta_line(base, d))}</div>
{''.join(rows)}</body></html>"""


def to_pdf(base: str) -> Path:
    if not Path(CHROME).exists():
        raise RuntimeError("Google Chrome not found — export Word instead, "
                           "or install Chrome for PDF export.")
    out = _out_path(base, "pdf")
    with tempfile.TemporaryDirectory() as td:
        page = Path(td) / "t.html"
        page.write_text(to_html(base), encoding="utf-8")
        subprocess.run(
            [CHROME, "--headless", "--disable-gpu", "--no-pdf-header-footer",
             f"--print-to-pdf={out}", str(page)],
            check=True, capture_output=True, timeout=60)
    return out


def export(base: str, fmt: str) -> Path:
    if fmt == "docx":
        return to_docx(base)
    if fmt == "pdf":
        return to_pdf(base)
    raise ValueError(f"unknown export format {fmt}")
